from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
from dotenv import load_dotenv
# from app.routers import resume
from llamaapi import LlamaAPI
import uuid
import os

load_dotenv()
app = FastAPI()

app.add_middleware(
    CORSMiddleware, 
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/app/templates", StaticFiles(directory="app/templates"), name="app/templates")

TEMPLATE_DIR = "./app/templates/pdf"
GENERATED_RESUME_DIR = "./generated_resumes"

api_token = os.getenv("LLAMA_API_KEY")

llama = LlamaAPI(api_token)
class SuggestionRequest(BaseModel):
    typeOfSuggestion: str ## job or project
    description: str


@app.get("/")
def read_root():
    return {"message" : "welcome to the resume builder api"}

@app.get("/templates")
def get_templates():
    templates = [{"name": os.path.splitext(f)[0], "path": f"/app/templates"} for f in os.listdir(TEMPLATE_DIR)]
    return {"templates": templates}

@app.get("/template-preview/{filename}")
def get_template_preview(filename: str):
    filepath = os.path.join(TEMPLATE_DIR, filename)
    if os.path.exists(filepath):
        return FileResponse(filepath)
    return {"error": "File not found"}

@app.post("/ai-suggestions")
def ai_suggestions(data: SuggestionRequest):
    description = data.description
    typeOfSuggestion = data.typeOfSuggestion
    try:
        api_request_json = {
            "model": "llama3.1-70b",
            "messages": [
                {"role": "system", "content": "You are a resume assistant."},
                {"role": "user", "content": f"""
                    Improve this {typeOfSuggestion} description for a resume: {description}

                    Provide exactly 3 improved suggestions. Format them as follows:
                    1. [First Suggestion]
                    2. [Second Suggestion]
                    3. [Third Suggestion]
                                    """},
                                ]
        }

        response = llama.run(api_request_json)
        response_text = response.json()["choices"][0]["message"]["content"]

        suggestions = response_text.split("\n")
        clean_suggestions = [s.strip() for s in suggestions if s.strip().startswith(("1.", "2.", "3."))]

        return {"suggestions": clean_suggestions}
    except Exception as e:
        return {"error": str(e)}
    
@app.post("/generate-resume")
async def generate_resume(request: Request):
    data = await request.json()
    personal_info = data["personalInfo"]
    work_experiences = data["workExperiences"]
    educations = data["educations"]
    skills = data["skills"]
    projects = data["projects"]
    template_name = data["template"]
    # Load the template
    template_path = f"./templates/docx/{template_name}.docx"
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Template not found")

    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("{Your Name}", personal_info["fullName"])
        paragraph.text = paragraph.text.replace("{Your Email}", personal_info["email"])
        paragraph.text = paragraph.text.replace("{Your Phone}", personal_info["phone"])
    doc = Document(template_path)

    # Fill in personal info
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("{Your Name}", personal_info["fullName"])
        paragraph.text = paragraph.text.replace("{Your Email}", personal_info["email"])
        paragraph.text = paragraph.text.replace("{Your Phone}", personal_info["phone"])
    
    # Fill in work experiences
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{Company}" in cell.text:
                    # Clear placeholder content
                    cell.text = ""
                    # Add all work experiences
                    for exp in work_experiences:
                        cell.add_paragraph(
                            f"{exp['Company']}, {exp['Location']} — {exp['Job Title']}\n"
                            f"{exp['Duration']}\n"
                            f"{exp['Job Description']}"
                        )

                if "{Institution}" in cell.text:
                    # Clear placeholder content
                    cell.text = ""
                    # Add all education entries
                    for edu in educations:
                        cell.add_paragraph(
                            f"{edu['Institution']}, {edu['Location']} — {edu['Degree']}\n"
                            f"{edu['Duration']}"
                        )

                if "{Project Name}" in cell.text:
                    # Clear placeholder content
                    cell.text = ""
                    # Add all projects
                    for proj in projects:
                        cell.add_paragraph(
                            f"{proj['Project Name']} — {proj['Detail']}\n"
                            f"{proj['Description']}"
                        )

    # Fill in skills (optional, as they can be added directly as text)
    for paragraph in doc.paragraphs:
        if "{Skill}" in paragraph.text:
            paragraph.text = ""
            for skill in skills:
                paragraph.add_run(f"- {skill}\n")

    # Save the filled template
    file_id = str(uuid.uuid4())
    output_path = f"{GENERATED_RESUME_DIR}/resume_{file_id}.docx"
    doc.save(output_path)

    return {"download_url": f"http://127.0.0.1:8000/download/{file_id}"}